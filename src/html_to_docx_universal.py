from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
import os
import base64
from io import BytesIO

def hex_to_rgb(hex_color):
    """Converte cor hexadecimal para RGB."""
    if hex_color.startswith('#'):
        hex_color = hex_color[1:]
    if len(hex_color) == 3:
        hex_color = ''.join([c*2 for c in hex_color])
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except:
        return (0, 0, 0)  # Preto como fallback

def parse_css_color(color_value):
    """Parse de cores CSS em diferentes formatos."""
    if not color_value:
        return None
        
    color_value = color_value.strip().lower()
    
    # Cores nomeadas comuns
    color_names = {
        'red': (255, 0, 0), 'green': (0, 128, 0), 'blue': (0, 0, 255),
        'black': (0, 0, 0), 'white': (255, 255, 255), 'gray': (128, 128, 128),
        'grey': (128, 128, 128), 'yellow': (255, 255, 0), 'orange': (255, 165, 0),
        'purple': (128, 0, 128), 'pink': (255, 192, 203), 'brown': (165, 42, 42),
        'navy': (0, 0, 128), 'teal': (0, 128, 128), 'lime': (0, 255, 0),
        'cyan': (0, 255, 255), 'magenta': (255, 0, 255), 'silver': (192, 192, 192),
        'maroon': (128, 0, 0), 'olive': (128, 128, 0)
    }
    
    if color_value in color_names:
        return color_names[color_value]
    
    # Cor hexadecimal
    if color_value.startswith('#'):
        return hex_to_rgb(color_value)
    
    # RGB
    rgb_match = re.match(r'rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_value)
    if rgb_match:
        return tuple(int(x) for x in rgb_match.groups())
    
    return None

def parse_font_size(size_value):
    """Parse de tamanho de fonte CSS."""
    if not size_value:
        return None
        
    size_value = size_value.strip().lower()
    
    # Tamanhos nomeados
    size_names = {
        'xx-small': 8, 'x-small': 10, 'small': 12, 'medium': 14,
        'large': 16, 'x-large': 18, 'xx-large': 24
    }
    
    if size_value in size_names:
        return size_names[size_value]
    
    # Pixels
    px_match = re.match(r'(\d+(?:\.\d+)?)px', size_value)
    if px_match:
        return int(float(px_match.group(1)))
    
    # Pontos
    pt_match = re.match(r'(\d+(?:\.\d+)?)pt', size_value)
    if pt_match:
        return int(float(pt_match.group(1)))
    
    # Em (aproximação)
    em_match = re.match(r'(\d+(?:\.\d+)?)em', size_value)
    if em_match:
        return int(float(em_match.group(1)) * 14)  # 14pt como base
    
    return None

def extract_inline_styles(element):
    """Extrai estilos inline de um elemento."""
    styles = {}
    style_attr = element.get('style', '')
    
    if style_attr:
        # Parse dos estilos inline
        style_pairs = style_attr.split(';')
        for pair in style_pairs:
            if ':' in pair:
                key, value = pair.split(':', 1)
                styles[key.strip().lower()] = value.strip()
    
    return styles

def extract_comprehensive_css_styles(soup):
    """Extrai estilos CSS de forma mais abrangente."""
    styles = {}
    
    # Estilos padrão
    default_styles = {
        'h1': {'font-size': '24pt', 'font-weight': 'bold'},
        'h2': {'font-size': '20pt', 'font-weight': 'bold'},
        'h3': {'font-size': '18pt', 'font-weight': 'bold'},
        'h4': {'font-size': '16pt', 'font-weight': 'bold'},
        'h5': {'font-size': '14pt', 'font-weight': 'bold'},
        'h6': {'font-size': '12pt', 'font-weight': 'bold'},
        'strong': {'font-weight': 'bold'},
        'b': {'font-weight': 'bold'},
        'em': {'font-style': 'italic'},
        'i': {'font-style': 'italic'},
        'u': {'text-decoration': 'underline'}
    }
    
    styles.update(default_styles)
    
    # Procurar tags style
    style_tags = soup.find_all('style')
    for style_tag in style_tags:
        css_content = style_tag.string
        if css_content:
            # Parse básico de CSS
            css_rules = re.findall(r'([^{]+)\{([^}]+)\}', css_content, re.DOTALL)
            
            for selector, properties in css_rules:
                selector = selector.strip()
                
                # Remover comentários
                properties = re.sub(r'/\*.*?\*/', '', properties, flags=re.DOTALL)
                
                # Parse das propriedades
                style_dict = {}
                props = properties.split(';')
                for prop in props:
                    if ':' in prop:
                        key, value = prop.split(':', 1)
                        style_dict[key.strip().lower()] = value.strip()
                
                # Armazenar estilos por seletor
                if selector.startswith('.'):
                    class_name = selector[1:].strip()
                    styles[f'class_{class_name}'] = style_dict
                elif selector.startswith('#'):
                    id_name = selector[1:].strip()
                    styles[f'id_{id_name}'] = style_dict
                else:
                    # Tag selector
                    tag_name = selector.strip().lower()
                    if tag_name in styles:
                        styles[tag_name].update(style_dict)
                    else:
                        styles[tag_name] = style_dict
    
    return styles

def apply_comprehensive_styles(element, run, paragraph, styles):
    """Aplica estilos de forma abrangente a um run."""
    if not run:
        return
    
    # Coletar todos os estilos aplicáveis
    applicable_styles = {}
    
    # 1. Estilos da tag
    tag_name = element.name.lower() if element.name else ''
    if tag_name and tag_name in styles:
        applicable_styles.update(styles[tag_name])
    
    # 2. Estilos de classe
    classes = element.get('class', [])
    for class_name in classes:
        class_key = f'class_{class_name}'
        if class_key in styles:
            applicable_styles.update(styles[class_key])
    
    # 3. Estilos de ID
    element_id = element.get('id')
    if element_id:
        id_key = f'id_{element_id}'
        if id_key in styles:
            applicable_styles.update(styles[id_key])
    
    # 4. Estilos inline (maior prioridade)
    inline_styles = extract_inline_styles(element)
    applicable_styles.update(inline_styles)
    
    # Aplicar estilos ao run
    
    # Cor do texto
    if 'color' in applicable_styles:
        color_rgb = parse_css_color(applicable_styles['color'])
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
    
    # Cor de fundo (limitado no DOCX)
    if 'background-color' in applicable_styles:
        bg_color_rgb = parse_css_color(applicable_styles['background-color'])
        if bg_color_rgb:
            # DOCX tem suporte limitado para cor de fundo de texto
            try:
                run.font.highlight_color = RGBColor(*bg_color_rgb)
            except:
                pass
    
    # Tamanho da fonte
    if 'font-size' in applicable_styles:
        font_size = parse_font_size(applicable_styles['font-size'])
        if font_size:
            run.font.size = Pt(font_size)
    
    # Família da fonte
    if 'font-family' in applicable_styles:
        font_family = applicable_styles['font-family'].strip('"\'')
        # Remover fallbacks
        font_family = font_family.split(',')[0].strip()
        run.font.name = font_family
    
    # Peso da fonte (negrito)
    if 'font-weight' in applicable_styles:
        weight = applicable_styles['font-weight'].lower()
        if weight in ['bold', 'bolder', '700', '800', '900']:
            run.bold = True
    
    # Estilo da fonte (itálico)
    if 'font-style' in applicable_styles:
        style = applicable_styles['font-style'].lower()
        if style in ['italic', 'oblique']:
            run.italic = True
    
    # Decoração do texto
    if 'text-decoration' in applicable_styles:
        decoration = applicable_styles['text-decoration'].lower()
        if 'underline' in decoration:
            run.underline = True
    
    # Alinhamento do parágrafo
    if paragraph and 'text-align' in applicable_styles:
        align = applicable_styles['text-align'].lower()
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif align == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def process_element_universal(element, doc, styles, parent_paragraph=None):
    """Processa elementos HTML de forma universal."""
    if element.name is None:
        # Texto simples
        text = str(element).strip()
        if text and parent_paragraph:
            run = parent_paragraph.add_run(text)
            return
        elif text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            return
    
    if not element.name:
        return
        
    tag_name = element.name.lower()
    
    # Elementos de bloco que criam novos parágrafos
    block_elements = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'blockquote', 'pre']
    
    if tag_name in block_elements:
        # Criar novo parágrafo
        paragraph = doc.add_paragraph()
        
        # Processar conteúdo do elemento
        for child in element.children:
            if child.name is None:
                # Texto simples
                text = str(child).strip()
                if text:
                    run = paragraph.add_run(text)
                    apply_comprehensive_styles(element, run, paragraph, styles)
            else:
                # Elemento filho
                if child.name.lower() in ['span', 'strong', 'b', 'em', 'i', 'u', 'a']:
                    # Elementos inline
                    text = child.get_text()
                    if text:
                        run = paragraph.add_run(text)
                        apply_comprehensive_styles(child, run, paragraph, styles)
                else:
                    # Elemento de bloco aninhado
                    process_element_universal(child, doc, styles)
    
    elif tag_name == 'table':
        process_table_universal(element, doc, styles)
    
    elif tag_name == 'ul' or tag_name == 'ol':
        process_list_universal(element, doc, styles)
    
    elif tag_name == 'br':
        if parent_paragraph:
            parent_paragraph.add_run().add_break()
        else:
            doc.add_paragraph()
    
    elif tag_name == 'img':
        process_image_universal(element, doc, parent_paragraph)
    
    elif tag_name in ['span', 'strong', 'b', 'em', 'i', 'u', 'a']:
        # Elementos inline
        if parent_paragraph:
            text = element.get_text()
            if text:
                run = parent_paragraph.add_run(text)
                apply_comprehensive_styles(element, run, parent_paragraph, styles)
        else:
            # Criar parágrafo se não há pai
            paragraph = doc.add_paragraph()
            text = element.get_text()
            if text:
                run = paragraph.add_run(text)
                apply_comprehensive_styles(element, run, paragraph, styles)
    
    else:
        # Outros elementos - processar filhos
        for child in element.children:
            process_element_universal(child, doc, styles, parent_paragraph)

def process_table_universal(table_element, doc, styles):
    """Processa tabelas de forma universal."""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Contar colunas
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    if max_cols == 0:
        return
    
    # Criar tabela no Word
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # Preencher tabela
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                word_cell = table.cell(i, j)
                
                # Limpar célula
                word_cell.text = ''
                
                # Processar conteúdo da célula
                cell_paragraph = word_cell.paragraphs[0]
                
                for child in cell.children:
                    if child.name is None:
                        text = str(child).strip()
                        if text:
                            run = cell_paragraph.add_run(text)
                            apply_comprehensive_styles(cell, run, cell_paragraph, styles)
                    else:
                        if child.name.lower() in ['span', 'strong', 'b', 'em', 'i', 'u']:
                            text = child.get_text()
                            if text:
                                run = cell_paragraph.add_run(text)
                                apply_comprehensive_styles(child, run, cell_paragraph, styles)
                        else:
                            text = child.get_text()
                            if text:
                                run = cell_paragraph.add_run(text)
                                apply_comprehensive_styles(child, run, cell_paragraph, styles)
                
                # Aplicar estilos da célula
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def process_list_universal(list_element, doc, styles):
    """Processa listas de forma universal."""
    items = list_element.find_all('li', recursive=False)
    
    for item in items:
        paragraph = doc.add_paragraph()
        
        # Adicionar marcador
        if list_element.name == 'ul':
            paragraph.style = 'List Bullet'
        else:
            paragraph.style = 'List Number'
        
        # Processar conteúdo do item
        for child in item.children:
            if child.name is None:
                text = str(child).strip()
                if text:
                    run = paragraph.add_run(text)
                    apply_comprehensive_styles(item, run, paragraph, styles)
            else:
                if child.name.lower() in ['span', 'strong', 'b', 'em', 'i', 'u']:
                    text = child.get_text()
                    if text:
                        run = paragraph.add_run(text)
                        apply_comprehensive_styles(child, run, paragraph, styles)
                else:
                    text = child.get_text()
                    if text:
                        run = paragraph.add_run(text)
                        apply_comprehensive_styles(child, run, paragraph, styles)

def process_image_universal(img_element, doc, parent_paragraph=None):
    """Processa imagens de forma universal."""
    src = img_element.get('src', '')
    
    if src.startswith('data:image'):
        # Imagem base64
        try:
            # Extrair dados da imagem
            header, data = src.split(',', 1)
            image_data = base64.b64decode(data)
            
            # Adicionar imagem ao documento
            if parent_paragraph:
                run = parent_paragraph.add_run()
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
            
            # Adicionar imagem (tamanho limitado)
            run.add_picture(BytesIO(image_data), width=Inches(4))
            
        except Exception as e:
            # Se falhar, adicionar texto alternativo
            alt_text = img_element.get('alt', '[Imagem]')
            if parent_paragraph:
                parent_paragraph.add_run(f'[{alt_text}]')
            else:
                doc.add_paragraph(f'[{alt_text}]')
    
    elif src.startswith('http'):
        # Imagem externa - adicionar apenas texto alternativo
        alt_text = img_element.get('alt', '[Imagem externa]')
        if parent_paragraph:
            parent_paragraph.add_run(f'[{alt_text}]')
        else:
            doc.add_paragraph(f'[{alt_text}]')
    
    elif os.path.exists(src):
        # Imagem local
        try:
            if parent_paragraph:
                run = parent_paragraph.add_run()
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
            
            run.add_picture(src, width=Inches(4))
        except:
            alt_text = img_element.get('alt', '[Imagem]')
            if parent_paragraph:
                parent_paragraph.add_run(f'[{alt_text}]')
            else:
                doc.add_paragraph(f'[{alt_text}]')

def convert_html_to_docx_universal(html_content, output_path):
    """Converte HTML para DOCX de forma universal."""
    try:
        # Parse do HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Criar documento Word
        doc = Document()
        
        # Extrair estilos CSS
        styles = extract_comprehensive_css_styles(soup)
        
        # Processar o body do HTML
        body = soup.find('body')
        if body:
            for child in body.children:
                process_element_universal(child, doc, styles)
        else:
            # Se não há body, processar todo o conteúdo
            for child in soup.children:
                if hasattr(child, 'name') and child.name:
                    process_element_universal(child, doc, styles)
        
        # Salvar documento
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Erro na conversão universal: {str(e)}")
        return False

if __name__ == "__main__":
    # Teste da função
    test_html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            .titulo { color: blue; font-size: 24px; text-align: center; font-weight: bold; }
            .subtitulo { color: red; font-size: 18px; font-weight: bold; }
            .texto-verde { color: green; font-family: Arial; }
            .destaque { background-color: yellow; color: black; font-weight: bold; }
        </style>
    </head>
    <body>
        <h1 class="titulo">Título Principal</h1>
        <h2 class="subtitulo">Subtítulo</h2>
        <p class="texto-verde">Este é um parágrafo em verde com fonte Arial.</p>
        <p>Texto normal com <span class="destaque">texto destacado</span> no meio.</p>
        <p style="color: purple; font-size: 16px;">Parágrafo com estilo inline roxo.</p>
    </body>
    </html>
    """
    
    convert_html_to_docx_universal(test_html, "teste_universal.docx")
    print("Conversão universal concluída!")