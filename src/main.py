from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, BackgroundTasks, Header, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from starlette.background import BackgroundTask
from pydantic import BaseModel, HttpUrl
import aiofiles
import requests
import tempfile
import os
from typing import Optional, Annotated
from file_converter import FileConverter
import logging
import asyncio
import base64
import subprocess
import re
import html
import json
import sys
import shutil
import uuid
import time
from datetime import datetime, timedelta
import threading
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from html_to_docx_universal import convert_html_to_docx_universal

# Função para verificar dependências críticas
def check_dependencies():
    """Verifica se todas as dependências críticas estão instaladas"""
    missing_deps = []
    
    # Verificar Pandoc
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True, timeout=10)
        if result.returncode != 0:
            missing_deps.append("Pandoc (comando não executou corretamente)")
    except FileNotFoundError:
        missing_deps.append("Pandoc (comando não encontrado)")
    except subprocess.TimeoutExpired:
        missing_deps.append("Pandoc (timeout na verificação)")
    except Exception as e:
        missing_deps.append(f"Pandoc (erro: {str(e)})")
    
    # Verificar LibreOffice (opcional, mas recomendado)
    try:
        result = subprocess.run(['soffice', '--version'], capture_output=True, text=True, timeout=10)
        if result.returncode != 0:
            print("AVISO: LibreOffice não está disponível. Alguns formatos podem não funcionar.")
    except FileNotFoundError:
        print("AVISO: LibreOffice não encontrado. Alguns formatos podem não funcionar.")
    except Exception:
        print("AVISO: Erro ao verificar LibreOffice.")
    
    # Verificar dependências Python críticas
    python_deps = {
        'fastapi': 'fastapi',
        'uvicorn': 'uvicorn', 
        'aiofiles': 'aiofiles',
        'requests': 'requests',
        'beautifulsoup4': 'bs4',
        'lxml': 'lxml'
    }
    for package_name, import_name in python_deps.items():
        try:
            __import__(import_name)
        except ImportError:
            missing_deps.append(f"Python package: {package_name}")
    
    if missing_deps:
        print("ERRO: Dependências críticas não encontradas:")
        for dep in missing_deps:
            print(f"  - {dep}")
        print("\nPara instalar as dependências:")
        print("1. Pandoc: https://pandoc.org/installing.html")
        print("2. LibreOffice: https://www.libreoffice.org/download/")
        print("3. Python packages: pip install -r requirements.txt")
        sys.exit(1)
    
    print("✓ Todas as dependências críticas estão disponíveis")

# Configuração de autenticação
def get_api_key():
    # Primeiro tenta ler de arquivo (Docker secret)
    api_key_file = os.getenv("API_KEY_FILE")
    if api_key_file and os.path.exists(api_key_file):
        with open(api_key_file, 'r') as f:
            return f.read().strip()
    
    # Fallback para variável de ambiente
    return os.getenv("API_KEY", "default-api-key-change-me")

API_KEY = get_api_key()

# Configuração da duração dos arquivos temporários
TEMP_FILE_DURATION_MINUTES = int(os.getenv("TEMP_FILE_DURATION_MINUTES", "15"))

# Diretório para arquivos temporários hospedados
TEMP_FILES_DIR = os.path.join(tempfile.gettempdir(), "api_temp_files")
os.makedirs(TEMP_FILES_DIR, exist_ok=True)

# Dicionário para rastrear arquivos temporários
temp_files_registry = {}
temp_files_lock = threading.Lock()

# Funções para gerenciar arquivos temporários
def cleanup_expired_files():
    """Remove arquivos temporários expirados"""
    current_time = datetime.now()
    expired_files = []
    
    with temp_files_lock:
        for file_id, file_info in list(temp_files_registry.items()):
            if current_time > file_info['expires_at']:
                expired_files.append((file_id, file_info['file_path']))
                del temp_files_registry[file_id]
    
    # Remove arquivos do sistema de arquivos
    for file_id, file_path in expired_files:
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                print(f"Arquivo temporário removido: {file_id}")
        except Exception as e:
            print(f"Erro ao remover arquivo temporário {file_id}: {str(e)}")

def start_cleanup_scheduler():
    """Inicia o agendador de limpeza de arquivos"""
    def cleanup_loop():
        while True:
            time.sleep(60)  # Verifica a cada minuto
            cleanup_expired_files()
    
    cleanup_thread = threading.Thread(target=cleanup_loop, daemon=True)
    cleanup_thread.start()

def store_temp_file(file_path: str, filename: str, format: str) -> str:
    """Armazena um arquivo temporário e retorna o ID único"""
    file_id = str(uuid.uuid4())
    expires_at = datetime.now() + timedelta(minutes=TEMP_FILE_DURATION_MINUTES)
    
    # Criar nome único para o arquivo
    file_extension = format.lower()
    temp_filename = f"{file_id}.{file_extension}"
    temp_file_path = os.path.join(TEMP_FILES_DIR, temp_filename)
    
    # Copiar arquivo para diretório temporário
    shutil.copy2(file_path, temp_file_path)
    
    # Registrar arquivo
    with temp_files_lock:
        temp_files_registry[file_id] = {
            'file_path': temp_file_path,
            'filename': filename,
            'format': format,
            'created_at': datetime.now(),
            'expires_at': expires_at
        }
    
    return file_id

# Verificar dependências na inicialização
check_dependencies()

# Iniciar agendador de limpeza
start_cleanup_scheduler()

app = FastAPI(
    title="API de Conversão de Arquivos",
    description="API para extrair texto de diversos formatos de arquivo",
    version="1.0.0"
)

# Função para verificar a API Key
async def verify_api_key(x_api_key: Annotated[str, Header()]):
    if x_api_key != API_KEY:
        raise HTTPException(
            status_code=401,
            detail="API Key inválida"
        )
    return x_api_key

converter = FileConverter()

class URLRequest(BaseModel):
    url: HttpUrl
    filename: Optional[str] = None

class GenerateFileRequest(BaseModel):
    file: str  # HTML bruto ou encodado em base64
    format: str  # Formato de saída (docx, pdf, etc.)

@app.get("/")
async def root():
    return {
        "message": "DEVFY - API de Conversão de Arquivos",
        "authentication": "Requer header 'x-api-key' para endpoints protegidos",
        "endpoints": {
            "/convert/url": "POST - Converter arquivo via URL (protegido)",
            "/convert/file": "POST - Converter arquivo binário (protegido)",
            "/generate": "POST - Gerar arquivo a partir de HTML (protegido)",
            "/generate/url": "POST - Gerar arquivo e retornar URL temporária (protegido)",
            "/temp/{file_id}": "GET - Download de arquivo temporário",
            "/formats": "GET - Formatos suportados (protegido)"
        }
    }

@app.get("/formats")
async def get_supported_formats(api_key: str = Depends(verify_api_key)):
    """Retorna os formatos de arquivo suportados"""
    return {
        "supported_formats": [
            "DOCX", "DOC", "XML", "YML", "XLSX", "XLS", "CSV", 
            "PDF", "TXT", "PPTX", "PPT", "HTML", "HTM", 
            "ODT", "ODP", "ODS", "JSON"
        ]
    }

@app.post("/convert/url")
async def convert_from_url(request: URLRequest, api_key: str = Depends(verify_api_key)):
    """Converte arquivo a partir de uma URL"""
    try:
        # Download do arquivo
        response = requests.get(str(request.url), timeout=30)
        response.raise_for_status()
        
        # Determina o nome do arquivo
        if request.filename:
            filename = request.filename
        else:
            filename = str(request.url).split('/')[-1]
            if '.' not in filename:
                raise HTTPException(status_code=400, detail="Não foi possível determinar a extensão do arquivo")
        
        # Salva temporariamente o arquivo
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{filename}") as temp_file:
            temp_file.write(response.content)
            temp_path = temp_file.name
        
        try:
            # Converte o arquivo
            extracted_text = await converter.convert_file(temp_path, filename)
            
            return JSONResponse(content={
                "success": True,
                "filename": filename,
                "url": str(request.url),
                "extracted_text": extracted_text,
                "file_size": len(response.content)
            })
        
        finally:
            # Remove arquivo temporário
            if os.path.exists(temp_path):
                os.unlink(temp_path)
                
    except requests.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Erro ao baixar arquivo: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro na conversão: {str(e)}")

@app.post("/convert/file")
async def convert_from_file(file: UploadFile = File(...), api_key: str = Depends(verify_api_key)):
    """Converte arquivo enviado diretamente"""
    try:
        # Valida se o arquivo foi enviado
        if not file.filename:
            raise HTTPException(status_code=400, detail="Nome do arquivo é obrigatório")
        
        # Salva temporariamente o arquivo
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{file.filename}") as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Converte e limpa o texto
            raw_text = await converter.convert_file(temp_path, file.filename)
            cleaned_text = converter.clean_text(raw_text)
            
            return JSONResponse(content={
                "success": True,
                "filename": file.filename,
                "extracted_text": cleaned_text,
                "total_characters": len(cleaned_text),
                "file_size": len(content),
                "content_type": file.content_type
            })
        
        finally:
            # Remove arquivo temporário
            if os.path.exists(temp_path):
                os.unlink(temp_path)
                
    except Exception as e:
            raise HTTPException(status_code=500, detail=f"Erro na conversão: {str(e)}")

@app.get("/temp/{file_id}")
async def download_temp_file(file_id: str):
    """Download de arquivo temporário hospedado"""
    with temp_files_lock:
        if file_id not in temp_files_registry:
            raise HTTPException(status_code=404, detail="Arquivo não encontrado ou expirado")
        
        file_info = temp_files_registry[file_id]
        
        # Verificar se o arquivo ainda existe
        if not os.path.exists(file_info['file_path']):
            del temp_files_registry[file_id]
            raise HTTPException(status_code=404, detail="Arquivo não encontrado")
        
        # Verificar se não expirou
        if datetime.now() > file_info['expires_at']:
            try:
                os.unlink(file_info['file_path'])
            except:
                pass
            del temp_files_registry[file_id]
            raise HTTPException(status_code=410, detail="Arquivo expirado")
    
    return FileResponse(
        path=file_info['file_path'],
        filename=file_info['filename'],
        media_type='application/octet-stream'
    )

async def convert_html_to_docx_enhanced(html_path: str, output_dir: str) -> str:
    """
    Converte HTML para DOCX preservando estilos CSS de forma mais precisa usando conversão universal.
    
    Args:
        html_path: Caminho para o arquivo HTML
        output_dir: Diretório de saída
        
    Returns:
        str: Caminho para o arquivo DOCX gerado
    """
    try:
        # Ler o arquivo HTML
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Usar a nova função de conversão universal
        output_path = os.path.join(output_dir, 'output.docx')
        success = convert_html_to_docx_universal(html_content, output_path)
        
        if success:
            return output_path
        else:
            raise Exception("Conversão universal falhou")
        
    except Exception as e:
        print(f"Erro na conversão HTML para DOCX: {str(e)}")
        # Fallback para pandoc
        return await fallback_pandoc_conversion(html_path, output_dir, 'docx')

def extract_css_styles(soup):
    """
    Extrai estilos CSS do HTML e converte para formato utilizável.
    """
    styles = {
        'bold': {'font-weight': 'bold'},
        'center': {'text-align': 'center'},
        'right': {'text-align': 'right'},
        'page-break': {'page-break-after': 'always'}
    }
    
    # Procurar tag style
    style_tag = soup.find('style')
    if style_tag:
        css_content = style_tag.string
        if css_content:
            # Parse básico de CSS
            import re
            
            # Extrair regras CSS
            css_rules = re.findall(r'([^{]+)\{([^}]+)\}', css_content)
            
            for selector, properties in css_rules:
                selector = selector.strip()
                if selector.startswith('.'):
                    class_name = selector[1:]
                    styles[class_name] = {}
                    
                    # Parse das propriedades
                    props = properties.split(';')
                    for prop in props:
                        if ':' in prop:
                            key, value = prop.split(':', 1)
                            styles[class_name][key.strip()] = value.strip()
    
    return styles

def process_html_element(element, doc, styles, parent_run=None):
    """
    Processa recursivamente elementos HTML e adiciona ao documento Word.
    """
    if element.name is None:
        # Texto simples
        text = str(element).strip()
        if text and parent_run:
            parent_run.text += text
        elif text:
            p = doc.add_paragraph()
            run = p.add_run(text)
        return
    
    # Processar diferentes tipos de elementos
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        # Cabeçalhos
        level = int(element.name[1])
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.bold = True
        
        # Tamanhos de fonte para cabeçalhos
        font_sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10}
        run.font.size = Pt(font_sizes.get(level, 12))
        
        # Aplicar alinhamento se houver classe
        apply_element_styles(element, p, run, styles)
        
    elif element.name == 'p':
        # Parágrafos
        p = doc.add_paragraph()
        process_paragraph_content(element, p, styles)
        apply_element_styles(element, p, None, styles)
        
    elif element.name in ['ul', 'ol']:
        # Listas
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph()
            p.style = 'List Bullet' if element.name == 'ul' else 'List Number'
            process_paragraph_content(li, p, styles)
            
    elif element.name == 'table':
        # Tabelas
        process_table(element, doc, styles)
        
    elif element.name == 'blockquote':
        # Citações
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.italic = True
        p.paragraph_format.left_indent = Inches(0.5)
        
    elif element.name == 'pre':
        # Código pré-formatado
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
    elif element.name == 'div':
        # Divs - verificar se é quebra de página
        if 'page-break' in element.get('class', []):
            doc.add_page_break()
        else:
            # Processar conteúdo da div
            for child in element.children:
                process_html_element(child, doc, styles)
    else:
        # Outros elementos - processar filhos
        for child in element.children:
            process_html_element(child, doc, styles)

def process_paragraph_content(element, paragraph, styles):
    """
    Processa o conteúdo de um parágrafo, incluindo spans com formatação.
    """
    for child in element.children:
        if child.name is None:
            # Texto simples
            text = str(child).strip()
            if text:
                run = paragraph.add_run(text)
        elif child.name == 'span':
            # Spans com possível formatação
            run = paragraph.add_run(child.get_text())
            apply_span_styles(child, run, styles)
        elif child.name in ['strong', 'b']:
            # Texto em negrito
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ['em', 'i']:
            # Texto em itálico
            run = paragraph.add_run(child.get_text())
            run.italic = True
        else:
            # Outros elementos inline
            run = paragraph.add_run(child.get_text())

def apply_element_styles(element, paragraph, run, styles):
    """
    Aplica estilos CSS a elementos do documento.
    """
    classes = element.get('class', [])
    
    for class_name in classes:
        if class_name in styles:
            style_props = styles[class_name]
            
            # Aplicar alinhamento
            if 'text-align' in style_props:
                align = style_props['text-align']
                if align == 'center':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Aplicar negrito
            if run and 'font-weight' in style_props:
                if style_props['font-weight'] == 'bold':
                    run.bold = True

def apply_span_styles(span, run, styles):
    """
    Aplica estilos específicos a spans.
    """
    classes = span.get('class', [])
    
    for class_name in classes:
        if class_name in styles:
            style_props = styles[class_name]
            
            # Aplicar negrito
            if 'font-weight' in style_props and style_props['font-weight'] == 'bold':
                run.bold = True

def process_table(table_element, doc, styles):
    """
    Processa uma tabela HTML e adiciona ao documento Word.
    """
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Contar colunas
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    # Criar tabela no Word
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # Preencher tabela
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                word_cell = table.cell(i, j)
                word_cell.text = cell.get_text().strip()
                
                # Aplicar negrito para cabeçalhos
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

async def fallback_pandoc_conversion(html_path: str, output_dir: str, format: str) -> str:
    """
    Conversão de fallback usando pandoc com opções aprimoradas.
    """
    output_file = os.path.join(output_dir, f'output.{format}')
    
    cmd = [
        'pandoc', 
        html_path, 
        '-o', output_file,
        '--standalone',
        '--preserve-tabs'
    ]
    
    if format == 'docx':
        cmd.extend([
            '--reference-doc=/dev/null',  # Usar template padrão
            '--wrap=none'
        ])
    
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    
    if result.returncode != 0:
        raise Exception(f"Pandoc fallback error: {result.stderr or result.stdout}")
    
    return output_file

def sanitize_html_content(html_content: str) -> str:
    """
    Sanitiza e corrige problemas comuns de escape e lint no HTML.
    
    Args:
        html_content: Conteúdo HTML bruto
        
    Returns:
        str: HTML sanitizado e corrigido
    """
    try:
        # 1. Corrigir aspas simples escapadas incorretamente
        # Padrão: '\'' -> '
        html_content = re.sub(r"\\'\\'", "'", html_content)
        
        # 2. Corrigir aspas duplas escapadas desnecessariamente em CSS
        # Padrão: \"Times New Roman\" -> "Times New Roman"
        html_content = re.sub(r'\\"([^"]*)\\"', r'"\1"', html_content)
        
        # 3. Remover caracteres de controle inválidos
        # Remove caracteres de controle exceto \n, \r, \t
        html_content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', html_content)
        
        # 4. Normalizar quebras de linha
        html_content = html_content.replace('\r\n', '\n').replace('\r', '\n')
        
        # 5. Corrigir entidades HTML mal formadas
        # Decodificar entidades HTML válidas
        html_content = html.unescape(html_content)
        
        # 6. Corrigir problemas comuns de CSS
        # Remover espaços extras em propriedades CSS
        html_content = re.sub(r'(\w+)\s*:\s*([^;]+);', r'\1: \2;', html_content)
        
        # 7. Validar e corrigir estrutura básica de tags
        # Garantir que tags importantes estejam fechadas
        if '<html' in html_content and '</html>' not in html_content:
            html_content += '</html>'
        if '<body' in html_content and '</body>' not in html_content:
            html_content = html_content.replace('</html>', '</body></html>')
        if '<head' in html_content and '</head>' not in html_content:
            html_content = html_content.replace('<body', '</head><body')
        
        # 8. Remover comentários HTML malformados
        html_content = re.sub(r'<!--[^>]*-->', '', html_content, flags=re.DOTALL)
        
        return html_content
        
    except Exception as e:
        # Se houver erro na sanitização, retorna o conteúdo original
        print(f"Erro na sanitização do HTML: {str(e)}")
        return html_content

@app.post("/generate/url")
async def generate_file_url(
    generate_request: GenerateFileRequest,
    request: Request,
    api_key: str = Depends(verify_api_key)
):
    """
    Gera um arquivo no formato especificado a partir de HTML e retorna uma URL temporária.
    
    Este endpoint funciona de forma similar ao /generate, mas em vez de retornar o arquivo diretamente,
    hospeda o arquivo temporariamente e retorna uma URL para download.
    
    **Características:**
    - Arquivo fica disponível por tempo limitado (configurável via TEMP_FILE_DURATION_MINUTES)
    - URL de download: /temp/{file_id}
    - Limpeza automática de arquivos expirados
    - Mesmo suporte de formatos do endpoint /generate
    
    **Exemplo de resposta:**
    ```json
    {
        "success": true,
        "file_id": "123e4567-e89b-12d3-a456-426614174000",
        "download_url": "/temp/123e4567-e89b-12d3-a456-426614174000",
        "filename": "generated_document.docx",
        "format": "docx",
        "expires_at": "2024-01-01T15:30:00",
        "expires_in_minutes": 15
    }
    ```
    
    Args:
        request: Objeto contendo o HTML (bruto ou Base64) e o formato de saída
        api_key: API key para autenticação
    
    Returns:
        JSONResponse: Informações sobre o arquivo gerado e URL para download
        
    Raises:
        HTTPException 400: Formato não suportado ou dados inválidos
        HTTPException 500: Erro na conversão ou dependências não encontradas
    """
    # Validar formato de saída
    supported_formats = ['docx', 'pdf', 'odt', 'txt', 'rtf', 'html']
    if generate_request.format.lower() not in supported_formats:
        raise HTTPException(
            status_code=400, 
            detail=f"Formato '{generate_request.format}' não suportado. Formatos suportados: {', '.join(supported_formats)}"
        )
    
    temp_html_path = None
    generated_file = None
    temp_dir = None
    
    try:
        # Decodificar o HTML se estiver em Base64; caso contrário, usar o texto diretamente
        try:
            html_content = base64.b64decode(generate_request.file).decode('utf-8')
        except Exception:
            html_content = generate_request.file  # Assume que o conteúdo já é HTML bruto
        
        # Sanitizar e corrigir problemas de escape e lint automaticamente
        html_content = sanitize_html_content(html_content)
        print(f"HTML sanitizado com sucesso. Tamanho: {len(html_content)} caracteres")
        
        # Criar arquivo HTML temporário
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
            temp_html.write(html_content)
            temp_html_path = temp_html.name
        
        # Gerar arquivo no formato especificado
        output_format = generate_request.format.lower()
        
        # Criar diretório temporário para saída
        temp_dir = tempfile.mkdtemp()
        
        if output_format in ['txt', 'docx', 'pdf', 'odt']:
            # Usar pandoc para conversões de HTML
            generated_file = os.path.join(temp_dir, f"output.{output_format}")
            
            cmd = ['pandoc', temp_html_path, '-o', generated_file]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            
            if result.returncode != 0:
                error_msg = f"Pandoc error: {result.stderr or result.stdout or 'Erro desconhecido'}"
                raise HTTPException(status_code=500, detail=error_msg)
            
        else:
            # Usar LibreOffice para outros formatos
            cmd = [
                'soffice',
                '--headless',
                '--invisible',
                '--nodefault',
                '--nolockcheck',
                '--nologo',
                '--norestore',
                '--convert-to', output_format,
                '--outdir', temp_dir,
                temp_html_path
            ]
            
            # Configurar ambiente
            env = os.environ.copy()
            env['HOME'] = '/tmp'
            env['TMPDIR'] = '/tmp'
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)
            
            if result.returncode != 0:
                error_msg = f"LibreOffice error: {result.stderr or result.stdout or 'Erro desconhecido'}"
                raise HTTPException(status_code=500, detail=error_msg)
        
        # Encontrar o arquivo gerado
        if output_format in ['txt', 'docx', 'pdf', 'odt']:
            # Para pandoc, o arquivo tem nome específico
            generated_file = os.path.join(temp_dir, f"output.{output_format}")
            if not os.path.exists(generated_file):
                files_in_dir = os.listdir(temp_dir)
                raise HTTPException(status_code=500, detail=f"Arquivo não foi gerado pelo pandoc. Arquivos no diretório: {files_in_dir}")
        else:
            # Para LibreOffice, procurar arquivo com extensão correta
            files_in_dir = os.listdir(temp_dir)
            output_files = [f for f in files_in_dir if f.endswith(f'.{output_format}')]
            
            if not output_files:
                raise HTTPException(status_code=500, detail=f"Arquivo não foi gerado. Arquivos no diretório: {files_in_dir}")
            
            generated_file = os.path.join(temp_dir, output_files[0])
        
        # Armazenar arquivo temporariamente e obter ID
        filename = f"generated_document.{output_format}"
        file_id = store_temp_file(generated_file, filename, output_format)
        
        # Calcular tempo de expiração
        expires_at = datetime.now() + timedelta(minutes=TEMP_FILE_DURATION_MINUTES)
        
        # Limpeza de arquivos temporários de processamento
        try:
            if temp_html_path and os.path.exists(temp_html_path):
                os.unlink(temp_html_path)
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except Exception:
            pass  # Ignora erros de limpeza
        
        # Construir URL completa baseada no host da requisição
        base_url = f"{request.url.scheme}://{request.url.netloc}"
        download_url = f"{base_url}/temp/{file_id}"
        
        return JSONResponse(content={
            "success": True,
            "file_id": file_id,
            "download_url": download_url,
            "filename": filename,
            "format": output_format,
            "expires_at": expires_at.isoformat(),
            "expires_in_minutes": TEMP_FILE_DURATION_MINUTES
        })
        
    except subprocess.TimeoutExpired:
        # Limpeza em caso de timeout
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail="Timeout na conversão do arquivo")
    except HTTPException:
        # Limpeza em caso de HTTPException
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise
    except Exception as e:
        # Limpeza em caso de erro geral
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Erro na geração do arquivo: {str(e)}")

@app.post("/generate")
async def generate_file(
    request: GenerateFileRequest,
    api_key: str = Depends(verify_api_key)
):
    """
    Gera um arquivo no formato especificado a partir de HTML.
    
    Este endpoint converte conteúdo HTML para diversos formatos de documento usando Pandoc e LibreOffice.
    
    **Formatos de entrada suportados:**
    - HTML bruto (texto HTML diretamente no campo 'file')
    - HTML codificado em Base64 (para compatibilidade com versões anteriores)
    
    **Formatos de saída suportados:**
    - **docx**: Microsoft Word Document (.docx) - via Pandoc
    - **pdf**: Portable Document Format (.pdf) - via Pandoc  
    - **txt**: Arquivo de texto simples (.txt) - via Pandoc
    - **odt**: OpenDocument Text (.odt) - via Pandoc
    - **rtf**: Rich Text Format (.rtf) - via LibreOffice
    - **html**: HTML sanitizado (.html) - processamento direto
    
    **Recursos adicionais:**
    - Sanitização automática de HTML (correção de escape, entidades, estrutura)
    - Detecção automática de formato de entrada (Base64 vs HTML bruto)
    - Validação de dependências (Pandoc, LibreOffice)
    - Limpeza automática de arquivos temporários
    - Timeout de 120 segundos para conversões
    
    **Exemplo de uso com HTML bruto:**
    ```json
    {
        "file": "<html><head><title>Teste</title></head><body><h1>Título</h1><p>Conteúdo</p></body></html>",
        "format": "docx"
    }
    ```
    
    **Exemplo de uso com Base64:**
    ```json
    {
        "file": "PGh0bWw+PGhlYWQ+PHRpdGxlPlRlc3RlPC90aXRsZT48L2hlYWQ+PGJvZHk+PGgxPlTDrXR1bG88L2gxPjxwPkNvbnRlw7pkbzwvcD48L2JvZHk+PC9odG1sPg==",
        "format": "pdf"
    }
    ```
    
    Args:
        request: Objeto contendo o HTML (bruto ou Base64) e o formato de saída
        api_key: API key para autenticação
    
    Returns:
        FileResponse: Arquivo gerado no formato especificado
        
    Raises:
        HTTPException 400: Formato não suportado ou dados inválidos
        HTTPException 500: Erro na conversão ou dependências não encontradas
    """
    # Validar formato de saída
    supported_formats = ['docx', 'pdf', 'odt', 'txt', 'rtf', 'html']
    if request.format.lower() not in supported_formats:
        raise HTTPException(
            status_code=400, 
            detail=f"Formato '{request.format}' não suportado. Formatos suportados: {', '.join(supported_formats)}"
        )
    
    temp_html_path = None
    generated_file = None
    
    try:
        # Decodificar o HTML se estiver em Base64; caso contrário, usar o texto diretamente
        try:
            html_content = base64.b64decode(request.file).decode('utf-8')
        except Exception:
            html_content = request.file  # Assume que o conteúdo já é HTML bruto
        
        # Sanitizar e corrigir problemas de escape e lint automaticamente
        html_content = sanitize_html_content(html_content)
        print(f"HTML sanitizado com sucesso. Tamanho: {len(html_content)} caracteres")
        
        # Criar arquivo HTML temporário
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
            temp_html.write(html_content)
            temp_html_path = temp_html.name
        
        # Debug: verificar se o arquivo foi criado e seu conteúdo
        print(f"Arquivo HTML criado: {temp_html_path}")
        print(f"Arquivo existe: {os.path.exists(temp_html_path)}")
        if os.path.exists(temp_html_path):
            with open(temp_html_path, 'r', encoding='utf-8') as f:
                content_preview = f.read()[:200]
                print(f"Primeiros 200 caracteres: {content_preview}")
                print(f"Tamanho do arquivo: {os.path.getsize(temp_html_path)} bytes")
        
        # Gerar arquivo no formato especificado
        output_format = request.format.lower()
        
        # Criar diretório temporário para saída
        temp_dir = tempfile.mkdtemp()
        
        if output_format == 'docx':
            # Usar conversão aprimorada para DOCX com melhor preservação de estilos
            generated_file = await convert_html_to_docx_enhanced(temp_html_path, temp_dir)
            
        elif output_format in ['txt', 'pdf', 'odt']:
            # Usar pandoc para outras conversões de HTML
            generated_file = os.path.join(temp_dir, f"output.{output_format}")
            
            # Comando pandoc com opções aprimoradas para preservar formatação
            cmd = [
                'pandoc', 
                temp_html_path, 
                '-o', generated_file,
                '--standalone',  # Documento completo
                '--preserve-tabs',  # Preservar tabs
                '--wrap=none'  # Não quebrar linhas automaticamente
            ]
            
            # Adicionar opções específicas para PDF
            if output_format == 'pdf':
                cmd.extend([
                    '--pdf-engine=wkhtmltopdf',  # Engine que preserva melhor CSS
                    '--css', temp_html_path  # Usar CSS do próprio HTML
                ])
            
            # Debug: imprimir comando
            print(f"Pandoc command: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            
            # Debug: imprimir informações sobre o resultado
            print(f"Return code: {result.returncode}")
            print(f"STDOUT: {result.stdout}")
            print(f"STDERR: {result.stderr}")
            
            if result.returncode != 0:
                error_msg = f"Pandoc error: {result.stderr or result.stdout or 'Erro desconhecido'}"
                raise HTTPException(status_code=500, detail=error_msg)
            
        else:
            # Usar LibreOffice para outros formatos
            cmd = [
                'soffice',
                '--headless',
                '--invisible',
                '--nodefault',
                '--nolockcheck',
                '--nologo',
                '--norestore',
                '--convert-to', output_format,
                '--outdir', temp_dir,
                temp_html_path
            ]
            
            # Configurar ambiente
            env = os.environ.copy()
            env['HOME'] = '/tmp'
            env['TMPDIR'] = '/tmp'
            
            # Debug: imprimir comando
            print(f"LibreOffice command: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)
            
            # Debug: imprimir informações sobre o resultado
            print(f"Return code: {result.returncode}")
            print(f"STDOUT: {result.stdout}")
            print(f"STDERR: {result.stderr}")
            
            if result.returncode != 0:
                error_msg = f"LibreOffice error: {result.stderr or result.stdout or 'Erro desconhecido'}"
                raise HTTPException(status_code=500, detail=error_msg)
        
        # Encontrar o arquivo gerado
        if output_format in ['txt', 'docx', 'pdf', 'odt']:
            # Para pandoc, o arquivo tem nome específico
            generated_file = os.path.join(temp_dir, f"output.{output_format}")
            if not os.path.exists(generated_file):
                files_in_dir = os.listdir(temp_dir)
                raise HTTPException(status_code=500, detail=f"Arquivo não foi gerado pelo pandoc. Arquivos no diretório: {files_in_dir}")
        else:
            # Para LibreOffice, procurar arquivo com extensão correta
            files_in_dir = os.listdir(temp_dir)
            output_files = [f for f in files_in_dir if f.endswith(f'.{output_format}')]
            
            if not output_files:
                raise HTTPException(status_code=500, detail=f"Arquivo não foi gerado. Arquivos no diretório: {files_in_dir}")
            
            generated_file = os.path.join(temp_dir, output_files[0])
        
        # Retornar o arquivo gerado
        filename = f"generated_document.{output_format}"
        
        def cleanup():
            """Função para limpeza de arquivos temporários"""
            try:
                if temp_html_path and os.path.exists(temp_html_path):
                    os.unlink(temp_html_path)
                if generated_file and os.path.exists(generated_file):
                    os.unlink(generated_file)
                if os.path.exists(temp_dir):
                    import shutil
                    shutil.rmtree(temp_dir)
            except Exception:
                pass  # Ignora erros de limpeza
        
        return FileResponse(
            path=generated_file,
            filename=filename,
            media_type='application/octet-stream',
            background=BackgroundTask(cleanup)
        )
        
    except subprocess.TimeoutExpired:
        # Limpeza em caso de timeout
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        raise HTTPException(status_code=500, detail="Timeout na conversão do arquivo")
    except HTTPException:
        # Limpeza em caso de HTTPException
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        raise
    except Exception as e:
        # Limpeza em caso de erro geral
        if temp_html_path and os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        raise HTTPException(status_code=500, detail=f"Erro na geração do arquivo: {str(e)}")

def cleanup_files(file_paths):
    """Remove arquivos temporários"""
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception:
            pass  # Ignora erros de limpeza

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)